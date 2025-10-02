
from docx import Document

doc = Document()
doc.add_heading('Laporan Penetration Testing – CVE-2012-2122', 0)

# Ringkasan Eksekutif
doc.add_heading('1. Ringkasan Eksekutif', level=1)
doc.add_paragraph(
    "Laporan ini menyajikan hasil pengujian keamanan terhadap layanan MySQL "
    "yang berjalan di host 10.33.102.225. Berdasarkan hasil identifikasi, enumerasi, dan eksploitasi, "
    "ditemukan bahwa versi MySQL yang digunakan (5.5.23) rentan terhadap CVE-2012-2122, yaitu kerentanan "
    "bypass autentikasi yang memungkinkan login tanpa kredensial valid."
)

# Deskripsi Kerentanan
doc.add_heading('2. Deskripsi Kerentanan – CVE-2012-2122', level=1)
doc.add_paragraph(
    "CVE ID: CVE-2012-2122\n"
    "Tipe: Authentication Bypass\n"
    "CVSS v2: 7.5 (High)\n"
    "Deskripsi:\n"
    "Kerentanan ini muncul akibat kegagalan fungsi perbandingan hash pada proses autentikasi MySQL. "
    "Dalam versi rentan, ada kemungkinan MySQL menerima password tidak valid sebagai sah jika terjadi kesalahan "
    "evaluasi dalam memcmp(). Rata-rata probabilitas keberhasilan adalah 1/256 pada tiap percobaan login.\n\n"
    "Produk Rentan:\n"
    "- MySQL < 5.1.63\n"
    "- MySQL < 5.5.24\n"
    "- MariaDB < 5.5.23\n\n"
    "Sumber Valid:\n"
    "- https://nvd.nist.gov/vuln/detail/CVE-2012-2122\n"
    "- https://www.exploit-db.com/exploits/19091"
)

# Metodologi Pengujian
doc.add_heading('3. Metodologi Pengujian', level=1)
doc.add_paragraph(
    "Tahap 1 – Scanning:\n"
    "Melakukan port dan versi detection menggunakan:\n"
    "nmap -sV -sC -Pn -O -oN nmap_results.txt 10.33.102.225\n\n"
    "Tahap 2 – Enumerasi:\n"
    "Mengekstrak hasil scanning untuk menemukan versi MySQL aktif.\n\n"
    "Tahap 3 – Analisis Versi:\n"
    "Dibandingkan dengan daftar versi rentan yang diketahui pada CVE-2012-2122.\n\n"
    "Tahap 4 – Eksploitasi:\n"
    "Melakukan brute-force login ke server MySQL menggunakan password acak dan mengirimkan perintah SHOW DATABASES;"
)

# Temuan
doc.add_heading('4. Temuan', level=1)
table = doc.add_table(rows=6, cols=2)
table.style = 'Table Grid'
data = [
    ("IP Target", "10.33.102.225"),
    ("Port Terbuka", "3306/tcp"),
    ("Service", "MySQL"),
    ("Versi", "5.5.23"),
    ("Status Login", "Berhasil tanpa password valid"),
    ("Jumlah Attempt", "409 percobaan"),
]

for i, row in enumerate(data):
    table.cell(i, 0).text = row[0]
    table.cell(i, 1).text = row[1]

doc.add_paragraph("Database Terdeteksi: information_schema, mysql, performance_schema, test")

# Analisis
doc.add_heading('5. Analisis', level=1)
doc.add_paragraph(
    "Target teridentifikasi menjalankan MySQL versi 5.5.23, yang sesuai dengan daftar versi rentan "
    "CVE-2012-2122. Hasil exploitasi menunjukkan bahwa attacker dapat memperoleh akses root hanya dengan brute-force acak, "
    "tanpa password valid, dan berhasil mengeksekusi perintah SQL."
)

# Dampak Potensial
doc.add_heading('6. Dampak Potensial', level=1)
doc.add_paragraph(
    "- Pengambilalihan penuh akses ke MySQL (dapat melihat, mengubah, atau menghapus data).\n"
    "- Peningkatan risiko lateral movement, terutama jika digunakan dalam sistem terintegrasi.\n"
    "- Pelanggaran integritas dan kerahasiaan data yang serius."
)

# Rekomendasi
doc.add_heading('7. Rekomendasi Remediasi', level=1)
doc.add_paragraph(
    "1. Segera upgrade MySQL ke versi aman:\n"
    "   - MySQL ≥ 5.5.24\n"
    "   - MariaDB ≥ 5.5.24\n"
    "2. Batasi akses remote ke port 3306 hanya dari alamat IP terpercaya.\n"
    "3. Audit log MySQL untuk mendeteksi akses mencurigakan.\n"
    "4. Terapkan firewall atau fail2ban untuk memblokir percobaan login brute-force.\n"
    "5. Gunakan autentikasi socket (unix_socket) jika memungkinkan pada sistem lokal."
)

# Referensi Teknis
doc.add_heading('8. Referensi Teknis', level=1)
doc.add_paragraph(
    "- CVE-2012-2122 – https://nvd.nist.gov/vuln/detail/CVE-2012-2122\n"
    "- Exploit PoC – https://www.exploit-db.com/exploits/19091\n"
    "- Oracle Patch Note – https://www.oracle.com/security-alerts/cpujul2012.html"
)

doc.save("Laporan_Pentesting_CVE-2012-2122.docx")
